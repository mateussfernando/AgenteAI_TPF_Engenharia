# app.py
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import re
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# --- Configurações iniciais ---
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    raise ValueError("A variável de ambiente GOOGLE_API_KEY não foi definida.")
genai.configure(api_key=API_KEY)

app = Flask(__name__)
CORS(app)  # Permite que o React acesse a API

PASTA_ARQUIVOS = "arquivos"
os.makedirs(PASTA_ARQUIVOS, exist_ok=True)

# --- Funções ---
INSTRUCAO_PERSONALIDADE = (
    "Você é um assistente de IA especializado em gerar relatórios chamada TPF - AI, planilhas e PDFs. "
    "Responda sempre em português, de forma clara, objetiva e profissional. "
    "Organize a resposta em tópicos numerados quando fizer sentido. "
    "Forneça apenas informações relevantes ao contexto de criação, edição e exportação de documentos."
)

def gerar_resposta(prompt: str) -> str:
    """Gera resposta usando Gemini com personalidade definida."""
    modelo = genai.GenerativeModel("gemini-1.5-flash")
    prompt_completo = INSTRUCAO_PERSONALIDADE + "\n\n" + prompt
    resposta = modelo.generate_content(prompt_completo)
    return resposta.text

def extrair_topicos(texto: str) -> list:
    """Extrai tópicos numerados da resposta."""
    topicos_encontrados = re.findall(r'^\s*(\d+\.\s*.*?)(?=\s*\d+\.|\Z)', texto, re.DOTALL | re.MULTILINE)
    if not topicos_encontrados:
        return [linha.strip() for linha in texto.split("\n") if linha.strip()]
    return topicos_encontrados

def salvar_docx(topicos: list, nome_base="relatorio_ia"):
    doc = Document()
    doc.add_heading("Relatório de Inteligência Artificial", 0)
    for topico in topicos:
        doc.add_paragraph(topico)
    caminho = os.path.join(PASTA_ARQUIVOS, f"{nome_base}.docx")
    doc.save(caminho)
    return caminho

def salvar_xlsx(topicos: list, nome_base="relatorio_ia"):
    wb = Workbook()
    ws = wb.active
    ws.title = "Relatório IA"
    for i, topico in enumerate(topicos, start=1):
        ws[f"A{i}"] = topico
    caminho = os.path.join(PASTA_ARQUIVOS, f"{nome_base}.xlsx")
    wb.save(caminho)
    return caminho

def salvar_pdf(topicos: list, nome_base="relatorio_ia"):
    caminho = os.path.join(PASTA_ARQUIVOS, f"{nome_base}.pdf")
    doc = SimpleDocTemplate(caminho, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = [Paragraph("Relatório de Inteligência Artificial", styles['Heading1']), Spacer(1,12)]
    for topico in topicos:
        flowables.append(Paragraph(topico, styles['Normal']))
        flowables.append(Spacer(1,6))
    doc.build(flowables)
    return caminho

# --- Rotas ---
@app.route('/', methods=['GET'])
def index():
    return jsonify({"status": "API rodando"}), 200

@app.route('/gerar', methods=['POST'])
def gerar():
    dados = request.get_json()
    prompt = dados.get("pergunta", "")
    if not prompt.strip():
        return jsonify({"resposta": "❌ Nenhuma pergunta recebida.", "arquivos": None}), 400

    # Gera resposta com personalidade em português
    resposta_bruta = gerar_resposta(prompt)
    topicos = extrair_topicos(resposta_bruta)

    # Cria nome base a partir de palavras-chave do prompt
    stop_words = ["o","a","os","as","um","uma","uns","umas","de","do","da","dos","das",
                  "para","em","no","na","nos","nas","com","por","sobre","que","e","se","ou"]
    palavras_chave = [p.lower() for p in re.split(r'\W+', prompt) if p.lower() not in stop_words and len(p) > 2]
    nome_base = "relatorio_" + "_".join(palavras_chave[:3]) if palavras_chave else "relatorio_gerado"

    # Salva arquivos
    docx_path = salvar_docx(topicos, nome_base)
    xlsx_path = salvar_xlsx(topicos, nome_base)
    pdf_path = salvar_pdf(topicos, nome_base)

    arquivos = {
        "docx": f"/download/{os.path.basename(docx_path)}",
        "xlsx": f"/download/{os.path.basename(xlsx_path)}",
        "pdf": f"/download/{os.path.basename(pdf_path)}"
    }

    return jsonify({"resposta": resposta_bruta, "arquivos": arquivos})

@app.route('/download/<path:filename>')
def download_file(filename):
    caminho = os.path.join(PASTA_ARQUIVOS, filename)
    if os.path.exists(caminho):
        return send_file(caminho, as_attachment=True)
    return "Arquivo não encontrado.", 404

if __name__ == '__main__':
    app.run(debug=True)

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    print("🟢 Servidor Flask iniciado!")
    print(f"➡️ Rota principal: http://localhost:{port}/")
    print(f"➡️ Rota de geração: http://localhost:{port}/gerar")
    print(f"➡️ Rota de download de arquivos: http://localhost:{port}/download/<nome_do_arquivo>")
    app.run(host="0.0.0.0", port=port, debug=True)